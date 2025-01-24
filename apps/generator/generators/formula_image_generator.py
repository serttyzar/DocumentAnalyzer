import random
from lib.problem_generator import *
from docx import Document
import matplotlib.pyplot as plt
from io import BytesIO
from generators.generator import ContentGenerator
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from faker import Faker
from localization.locale_strategy import LocaleStrategy
from docx.shared import Pt

class FormulaImageGenerator(ContentGenerator):
    """Генератор формул в виде изображений"""


    def __init__(self, font_size, locale_strategy: LocaleStrategy,  width: float = 1, height: float = 1):
        self.font_size = font_size
        self.width = width
        self.height = height
        self.locale_strategy = locale_strategy
        self.problem_generator = ProblemGenerator()
        self.problem_generator.add_expander(l_sum)
        self.problem_generator.add_expander(l_div)
        self.problem_generator.add_expander(l_pow)
        self.problem_generator.add_expander(l_sqrt)
        self.problem_generator.add_expander(l_int)
        self.problem_generator.add_expander(l_sig)
        self.formula_count = 0

    def generate(self, document: Document):
        """Создает формулу и вставляет ее в документ как изображение"""
        self.formula_count += 1
        # Сгенерировать формулу
        value = random.randint(1, 100)
        depth = random.randint(2, 6)
        latex_expression = self.problem_generator.randexpr(value, depth)
        
        # Преобразовать формулу в изображение
        img_stream = self.latex_to_image(latex_expression)
        
        
        
        # Добавить изображение в документ
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(img_stream, width=Pt(random.randint(55*depth, 75*depth)), height = Pt(random.randint(20*depth, 30*depth)))
        # Добавить подпись в документ
        if random.choices([True, False], weights=[0.5, 0.5])[0]:
            self._add_caption(document)


    def latex_to_image(self, latex_input):
        """Конвертирует формулу в изображение и возвращает поток данных"""
        fig, ax = plt.subplots(figsize=(self.width, self.height))
        ax.text(0.5, 0.5, f"${latex_input}$", fontsize=random.randint(35,40), ha='center')
        ax.axis('off')
        
        buf = BytesIO()
        fig.savefig(buf, format='png', pad_inches=0, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    
    def _add_caption(self, document):
        fake = Faker(self.locale_strategy.get_locale())
        caption_text = random.choice([ 
            f"Формула {self.formula_count} - {fake.sentence()}", 
            f"Формула {self.formula_count}"
        ])

        caption_paragraph = document.add_paragraph(caption_text)
        caption_paragraph.alignment = random.choices([WD_PARAGRAPH_ALIGNMENT.CENTER, WD_PARAGRAPH_ALIGNMENT.LEFT, WD_PARAGRAPH_ALIGNMENT.RIGHT], weights=[0.9, 0.05, 0.05])[0]
        caption_paragraph.runs[0].font.size = Pt(self.font_size)