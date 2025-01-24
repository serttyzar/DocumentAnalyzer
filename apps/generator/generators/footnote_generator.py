from faker import Faker
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from generators.generator import ContentGenerator
from localization.locale_strategy import LocaleStrategy

class FootnoteGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy: LocaleStrategy):
        self.font_size = font_size
        self.locale_strategy = locale_strategy
        
    def generate(self, document):
        fake = Faker(self.locale_strategy.get_locale())
        footnote_text = fake.sentence()
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{fake.text(max_nb_chars=50)}[См. сноску]")
        
        footnote_paragraph = document.add_paragraph()
        footnote_paragraph.add_run(f"Сноска: {footnote_text}")
        footnote_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        footnote_paragraph.runs[0].font.size = Pt(8)