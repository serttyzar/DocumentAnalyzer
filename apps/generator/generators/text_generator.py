import random

from faker import Faker
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from generators.generator import ContentGenerator
from localization.locale_strategy import LocaleStrategy

class TextGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy):
        self.font_size = font_size
        self.locale_strategy: LocaleStrategy = locale_strategy
        
    def generate(self, document):
        if random.choice([True, False]):
            self._generate_columns(document)
        else:
            self._generate_single_paragraph(document)
    
    def _generate_single_paragraph(self, document):
        fake = Faker(self.locale_strategy.get_locale())
        text = fake.text(max_nb_chars=random.randint(500, 4000))

        paragraph = document.add_paragraph(text)
        paragraph.alignment = random.choice([WD_PARAGRAPH_ALIGNMENT.LEFT, 
                                        WD_PARAGRAPH_ALIGNMENT.CENTER,
                                        WD_PARAGRAPH_ALIGNMENT.RIGHT,
                                        WD_PARAGRAPH_ALIGNMENT.JUSTIFY])
        run = paragraph.runs[0]
        run.font.size = Pt(self.font_size)
        paragraph.paragraph_format.first_line_indent = Pt(25)
    
    def _generate_columns(self, document):
        num_rows = 1
        num_cols = random.choices([1, 2 ,3], weights=[0.2, 0.5, 0.3])[0]
        table = document.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Normal Table'
        color = 'FFFFFF'
        fake = Faker(self.locale_strategy.get_locale())
        # text: str = fake.text(max_nb_chars=random.randint(500, 1000))

        for row in table.rows:
                for cell in row.cells:
                    run = cell.paragraphs[0].add_run(fake.text(max_nb_chars=random.randint(300, 600)))
                    run.font.size = Pt(self.font_size)                 
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)))
                    cell.paragraphs[0].alignment = random.choice([
                    WD_PARAGRAPH_ALIGNMENT.LEFT,
                    WD_PARAGRAPH_ALIGNMENT.CENTER,
                    WD_PARAGRAPH_ALIGNMENT.RIGHT
                ])
        table.alignment = random.choice([
            WD_PARAGRAPH_ALIGNMENT.CENTER, 
            WD_PARAGRAPH_ALIGNMENT.LEFT, 
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        ])