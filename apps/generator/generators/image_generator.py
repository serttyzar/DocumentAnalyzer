import random
from faker import Faker
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from generators.generator import ContentGenerator
from lib.random_picture import RandomPictureGenerator
from localization.locale_strategy import LocaleStrategy


class ImageGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy: LocaleStrategy, picture_generator: RandomPictureGenerator):
        self.font_size = font_size
        self.locale_strategy = locale_strategy
        self.picture_generator = picture_generator
        self.image_count = 0
        
    def generate(self, document):
        """Добавляет изображение с подписью"""
        self.image_count += 1

        if not (add_caption_after:=random.choices([True, False], weights=[0.9, 0.1])[0]):
            self._add_caption(document)
        
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.add_run().add_picture(self.picture_generator.get_random_image(), width=Pt(random.randint(150, 400)))
        
        if add_caption_after:
            self._add_caption(document)
    
    def _add_caption(self, document):
        fake = Faker(self.locale_strategy.get_locale())
        caption_text = random.choice([
            f"Рис. {self.image_count}", 
            f"Рисунок {self.image_count} - {fake.sentence()}", 
            f"Рисунок {self.image_count}"
        ])
        
        # Настройки подписи
        caption_paragraph = document.add_paragraph(caption_text)
        caption_paragraph.alignment = random.choices([WD_PARAGRAPH_ALIGNMENT.CENTER, WD_PARAGRAPH_ALIGNMENT.LEFT, WD_PARAGRAPH_ALIGNMENT.RIGHT], weights=[0.9, 0.05, 0.05])[0]
        caption_paragraph.runs[0].font.size = Pt(self.font_size)