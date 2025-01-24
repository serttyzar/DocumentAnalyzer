import random
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from faker import Faker
from generators.generator import ContentGenerator
from localization.locale_strategy import LocaleStrategy

class HeaderFooterGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy: LocaleStrategy):
        self.font_size = font_size
        self.locale_strategy = locale_strategy
    
    def generate(self, document):
        section = document.sections[0]
        self._setup_header(section.header)
        self._setup_footer(section.footer)
        
    def _setup_header(self, header):
        """Настраивает верхний колонтитул"""
        fake = Faker(self.locale_strategy.get_locale())
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = fake.sentence(nb_words=random.randint(2, 4))
        self._format_paragraph(header_paragraph)
        
    def _setup_footer(self, footer):
        """Настраивает нижний колонтитул"""
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.add_run().add_break()
        self._add_page_number(footer_paragraph)
        self._format_paragraph(footer_paragraph)
        
    def _add_page_number(self, paragraph):
        """Добавляет номер текущей страницы в нижний колонтитул"""
        run = paragraph.add_run()
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        instrText = parse_xml(r'<w:instrText {} xml:space="preserve">PAGE</w:instrText>'.format(nsdecls('w')))
        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Центрируем номер страницы
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    def _format_paragraph(self, paragraph):
        """Форматирует параграф колонтитула"""
        for run in paragraph.runs:
            run.font.size = Pt(self.font_size)