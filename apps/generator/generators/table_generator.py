from pydoc import doc
import random
from xml.dom.minidom import DocumentType
from docx import Document
from docx.table import Table
from faker import Faker
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from mimesis import Datetime, Numeric, Text
from generators.generator import ContentGenerator
from localization.locale_strategy import LocaleStrategy


class TableGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy: LocaleStrategy):
        self.font_size = font_size
        self.locale_strategy = locale_strategy
        self.table_count = 0
        self.colors = [hex(random.randint(0, 16777215))[2:] for _ in range(100)]
        
    def generate(self, document: Document):
        self.table_styles = self._get_styles(document)
        self.table_count += 1
        self._add_caption(document)
        
        # Создание таблицы
        num_rows = random.randint(2, 5)
        num_cols = random.randint(2, 8)
        table: Table = document.add_table(rows=num_rows, cols=num_cols)
        
        self._apply_table_style(table) 
        
        self._fill_table(table)     
        
        # Выравнивание таблицы
        table.alignment = random.choice([
            WD_PARAGRAPH_ALIGNMENT.CENTER, 
            WD_PARAGRAPH_ALIGNMENT.LEFT, 
            WD_PARAGRAPH_ALIGNMENT.RIGHT 
        ])

        spacing_paragraph = document.add_paragraph()
        spacing_paragraph.paragraph_format.space_before = Pt(2)
        spacing_paragraph.paragraph_format.space_after = Pt(4)
        
    def _apply_table_style(self, table):
        """Применяет стиль к таблице"""
        if random.choice([True, False]):
            table.style = 'Table Grid'
        else:
            table.style = random.choice(self.table_styles)
            
    def _fill_table(self, table: Table):
        is_colored = random.choices([True, False], weights=[0.8, 0.2])[0]
        painting_style = random.choices([0, 1, 2], weights=[0.79, 0.2, 0.01])[0]
        
        cell_colors = self._get_cell_colors(is_colored)
        
        if painting_style == 0:
            self._fill_table_striped(table, cell_colors)
        elif painting_style == 1:
            self._fill_table_two_colors(table, cell_colors)
        else:
            self._fill_table_random_colors(table)
        
    def _fill_table_two_colors(self, table: Table, cell_colors):
        """Заполняет таблицу двумя цветами"""
        for row in table.rows:
            for cell in row.cells:
                self._fill_cell(cell)
                color = random.choices(cell_colors, weights=[0.7, 0.3])[0]
                self._set_cell_color(cell, color)
                
    def _fill_table_random_colors(self, table):
        """Заполняет таблицу случайными цветами"""
        for row in table.rows:
            for cell in row.cells:
                self._fill_cell(cell)
                self._set_cell_color(cell, random.choice(self.colors))        
                
    def _fill_table_striped(self, table: Table, cell_colors):
        """Заполняет таблицу по строкам"""
        row_check = 1
        for row in table.rows:
            row_check *= -1
            for cell in row.cells:
                self._fill_cell(cell)
                color = cell_colors[0] if row_check > 0 else cell_colors[1]
                self._set_cell_color(cell, color)
    
    def _fill_cell(self, cell):
        """Заполняет ячейку"""
        content_type = random.choice(['number', 'date', 'text'])
        cell.text = self._get_cell_content(content_type)
        
        # Настройка выравнивания текста в ячейке
        cell.paragraphs[0].alignment = random.choice([
            WD_PARAGRAPH_ALIGNMENT.LEFT,
            WD_PARAGRAPH_ALIGNMENT.CENTER,
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        ]) 

        # TODO: Проверить
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(self.font_size)
    
    def _set_cell_color(self, cell, color):
        """Устанавливает цвет фона ячейки""" 
        cell._element.get_or_add_tcPr().append(
            parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        )
         
    def _get_cell_colors(self, is_colored: bool):
        """Возвращает цвета для ячейки"""
        if not is_colored:
            return ["d3d3d3", "d3d3d3"]
        return [random.choice(self.colors), random.choice(self.colors)]
    
    def _get_cell_content(self, content_type):
        """Генерирует содержимое ячейки"""
        fake = Faker(self.locale_strategy.get_locale())
        
        if content_type == 'number':
            return str(Numeric().integers(start=1, end=1000000)[0])
        elif content_type == 'date':
            return Datetime().date().strftime("%Y-%m-%d")
        elif content_type == 'text':
            return Text(self.locale_strategy.get_locale()).word() if random.choice([True, False]) else fake.text(max_nb_chars=50)
        return "" 
    
    def _get_styles(self, document):
        styles = document.styles
        table = document.add_table(rows = 1,cols = 1)
        table_styles = []
        for s in styles:
            try:
                table.style = s
                table_styles.append(s)
            except:
                pass
        self._delete_table(document, -1)
        return table_styles
        
    def _delete_table(self, document, table):
        document.tables[table]._element.getparent().remove(document.tables[table]._element)
        
    def _add_caption(self, document):
        """Добавляет подпись к таблице"""
        fake = Faker(self.locale_strategy.get_locale())
        
        caption_text = random.choice([
            f"Табл. {self.table_count}", 
            f"Таблица {self.table_count} - {fake.sentence()}", 
            "Таблица."
        ])
        
        caption_paragraph = document.add_paragraph(caption_text)
        caption_paragraph.runs[0].font.size = Pt(self.font_size)
        caption_paragraph.alignment = random.choices(
            [WD_PARAGRAPH_ALIGNMENT.CENTER, WD_PARAGRAPH_ALIGNMENT.LEFT, WD_PARAGRAPH_ALIGNMENT.RIGHT], 
            weights=[0.9, 0.05, 0.05]
        )[0]