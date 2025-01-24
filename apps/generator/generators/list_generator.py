import random
from faker import Faker
from docx.shared import Pt
from generators.generator import ContentGenerator
from lib.paragraph_list import ParagraphExt
from localization.locale_strategy import LocaleStrategy


class ListGenerator(ContentGenerator):
    def __init__(self, font_size, locale_strategy):
        self.font_size = font_size
        self.locale_strategy: LocaleStrategy = locale_strategy
        
    def generate(self, document):
        fake = Faker(self.locale_strategy.get_locale())
        list_type = random.choice(['List Number', 'List Bullet'])
        font_size = self.font_size

        for i in range(random.randint(2, 8)):
            if list_type == 'List Number':
                list_paragraph = ParagraphExt(document.add_paragraph(style=list_type))
                if i == 0:
                    list_paragraph.restart_numbering()
            else:
                list_paragraph = document.add_paragraph(style=list_type)
            list_paragraph.style.font.size = Pt(font_size)
            list_paragraph.paragraph_format.left_indent = Pt(35)
            list_paragraph.paragraph_format.hanging_indent = Pt(12)
            list_paragraph.paragraph_format.space_after = Pt(4)
        
            item = list_paragraph.add_run(fake.text(max_nb_chars=random.randint(50, 300)))